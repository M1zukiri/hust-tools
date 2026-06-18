"""
DOCX/DOC 转 Markdown 转换器
使用 python-docx 库
"""

from pathlib import Path
from typing import Optional, Dict, Any
import re

from converter_base import BaseConverter, ConversionResult, ConversionStatus


class DocxConverter(BaseConverter):
    """Word文档转Markdown转换器"""
    
    supported_extensions = ['docx', 'doc']
    supported_mime_types = [
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'application/msword'
    ]
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        super().__init__(config)
        self.preserve_images = config.get('preserve_images', False) if config else False
        self.image_output_dir = config.get('image_output_dir', None) if config else None
    
    def convert(self, file_path: Path, **options) -> ConversionResult:
        """
        将DOCX文件转换为Markdown
        
        Args:
            file_path: DOCX文件路径
            **options: 额外选项
                - preserve_images: 是否保留图片
                - image_output_dir: 图片输出目录
                
        Returns:
            ConversionResult: 转换结果
        """
        if not self.validate_file(file_path):
            return self._create_error_result(file_path, "文件验证失败", "docx")
        
        try:
            from docx import Document
        except ImportError:
            return self._create_error_result(
                file_path, 
                "缺少依赖: python-docx。请运行: pip install python-docx",
                "docx"
            )
        
        try:
            doc = Document(str(file_path))
            md_content = self._document_to_markdown(doc, file_path, **options)
            
            metadata = self.extract_metadata(file_path)
            metadata['paragraph_count'] = len(doc.paragraphs)
            metadata['table_count'] = len(doc.tables)
            
            return ConversionResult(
                content=md_content,
                status=ConversionStatus.SUCCESS,
                source_format='docx',
                metadata=metadata,
                errors=[]
            )
            
        except Exception as e:
            self.logger.error(f"转换失败: {e}")
            return self._create_error_result(file_path, str(e), "docx")
    
    def _document_to_markdown(self, doc, file_path: Path, **options) -> str:
        """
        将Document对象转换为Markdown
        
        Args:
            doc: python-docx Document对象
            file_path: 文件路径
            **options: 额外选项
            
        Returns:
            str: Markdown内容
        """
        md_parts = []
        
        # 处理段落
        for para in doc.paragraphs:
            md_text = self._paragraph_to_markdown(para)
            if md_text.strip():
                md_parts.append(md_text)
        
        # 处理表格
        for table in doc.tables:
            md_table = self._table_to_markdown(table)
            if md_table.strip():
                md_parts.append(md_table)
        
        return '\n\n'.join(md_parts)
    
    def _paragraph_to_markdown(self, para) -> str:
        """
        将段落转换为Markdown
        
        Args:
            para: 段落对象
            
        Returns:
            str: Markdown文本
        """
        text = para.text.strip()
        if not text:
            return ""
        
        # 获取段落样式
        style_name = para.style.name.lower() if para.style else ""
        
        # 根据样式确定标题级别
        if 'heading 1' in style_name or para.style.name == 'Title':
            return f"# {text}"
        elif 'heading 2' in style_name:
            return f"## {text}"
        elif 'heading 3' in style_name:
            return f"### {text}"
        elif 'heading 4' in style_name:
            return f"#### {text}"
        elif 'heading 5' in style_name:
            return f"##### {text}"
        elif 'heading 6' in style_name:
            return f"###### {text}"
        
        # 处理列表
        if para._p is not None:
            # 检查是否是列表项
            numPr = para._p.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
            if numPr is not None:
                # 简单处理为无序列表
                return f"- {text}"
        
        # 处理粗体和斜体
        md_text = ""
        for run in para.runs:
            run_text = run.text
            if run.bold and run.italic:
                run_text = f"***{run_text}***"
            elif run.bold:
                run_text = f"**{run_text}**"
            elif run.italic:
                run_text = f"*{run_text}*"
            
            # 处理代码样式
            if run.font and run.font.name and 'code' in run.font.name.lower():
                run_text = f"`{run_text}`"
            
            md_text += run_text
        
        return md_text if md_text else text
    
    def _table_to_markdown(self, table) -> str:
        """
        将表格转换为Markdown
        
        Args:
            table: 表格对象
            
        Returns:
            str: Markdown表格
        """
        if not table.rows:
            return ""
        
        md_lines = []
        
        # 表头
        header_cells = [cell.text.strip() for cell in table.rows[0].cells]
        md_lines.append('| ' + ' | '.join(header_cells) + ' |')
        md_lines.append('|' + '|'.join(['---' for _ in header_cells]) + '|')
        
        # 数据行
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            md_lines.append('| ' + ' | '.join(cells) + ' |')
        
        return '\n'.join(md_lines)
